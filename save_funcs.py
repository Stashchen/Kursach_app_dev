from docx import Document
from docx.shared import Inches
import openpyxl
from openpyxl.styles import Font, Alignment


def save_word():
    doc = Document()
    doc.add_heading('Result', 0)
    doc.add_picture('images/photo.png', width=Inches(6))
    doc.add_paragraph('Created by Pavel Stashchenko, group 10701118')
    doc.save('files/word_result.docx')


def save_exel():
    wb = openpyxl.Workbook() # Создание проекта
    ws = wb.active # Активация одного листа в excel
    ws['A1'].font = Font(size=16)
    ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
    ws['A1'] = 'RESULT'
    ws.merge_cells('A1:I1')
    img = openpyxl.drawing.image.Image('images/photo.png')
    img.anchor = 'A2'
    ws.add_image(img)
    ws['A37'].alignment = Alignment(horizontal='center', vertical='center')
    ws['A37'] = 'Created by Pavel Stashchenko, group 10701118'
    ws.merge_cells('A37:I37')
    wb.save('files/excel_result.xlsx')


